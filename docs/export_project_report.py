#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

os.makedirs(os.path.dirname(__file__), exist_ok=True)

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

BLUE = RGBColor(0x1F, 0x49, 0x7D)
GRAY = RGBColor(0x60, 0x60, 0x60)

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)

def add_section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text, bold_part=None, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    if bold_part and bold_part in text:
        before, after = text.split(bold_part, 1)
        if before:
            r = p.add_run(before); r.font.size = Pt(10.5)
        r2 = p.add_run(bold_part); r2.bold = True; r2.font.size = Pt(10.5)
        if after:
            r3 = p.add_run(after); r3.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows, 1):
        row = t.rows[ri].cells
        for ci, val in enumerate(row_data):
            row[ci].text = str(val)
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 0:
                tc = row[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EEF2F7')
                tcPr.append(shd)
    doc.add_paragraph()

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════
add_title(doc, "RESPOND Guatemala — Trauma Registry")
add_subtitle(doc, "Platform Overview: Infrastructure, Costs, Storage & Deployment\nGenerated March 2026")

# ════════════════════════════════════════════════════════════
# 1. PLATFORM & TECH STACK
# ════════════════════════════════════════════════════════════
add_section(doc, "1. Platform & Tech Stack")
add_bullet(doc, "Framework: Next.js + TypeScript (React 19)", bold_part="Next.js + TypeScript")
add_bullet(doc, "Database & Authentication: Supabase (PostgreSQL cloud)", bold_part="Supabase (PostgreSQL cloud)")
add_bullet(doc, "Hosting: Vercel — live at trauma-registry.vercel.app", bold_part="Vercel")
add_bullet(doc, "Offline support: PWA with IndexedDB — syncs automatically when internet is available", bold_part="PWA with IndexedDB")
add_bullet(doc, "Languages: Spanish & English (bilingual interface)")
add_bullet(doc, "Works on: tablet, phone, and desktop browser")

# ════════════════════════════════════════════════════════════
# 2. WHAT THE APP DOES
# ════════════════════════════════════════════════════════════
add_section(doc, "2. What the App Does")
add_bullet(doc, "16-step data collection form per trauma patient", bold_part="16-step data collection form")
add_bullet(doc, "Captures: demographics, injury mechanism, pre-hospital care, vitals, GCS, body map, AIS/ISS scoring, procedures, surgery, and outcome", indent=1)
add_bullet(doc, "Auto-calculates: GCS total, Shock Index, ISS score, Length of Stay", bold_part="Auto-calculates:")
add_bullet(doc, "Interactive body map for injury localization", bold_part="Interactive body map")
add_bullet(doc, "Dashboard: mortality rate, injury mechanisms, monthly admission trends", bold_part="Dashboard:")
add_bullet(doc, "Data export to Excel or CSV with custom field selection", bold_part="Data export")
add_bullet(doc, "Role-based access: Registrar / Admin / Viewer", bold_part="Role-based access:")
add_bullet(doc, "Offline-first: records saved locally and synced to cloud when connected", bold_part="Offline-first:")

# ════════════════════════════════════════════════════════════
# 3. CURRENT INFRASTRUCTURE COSTS
# ════════════════════════════════════════════════════════════
add_section(doc, "3. Current Infrastructure Costs")

add_table(doc,
    headers=['Service', 'Purpose', 'Plan', 'Cost/Month'],
    rows=[
        ('Vercel',   'Web app hosting',          'Hobby (Free)', '$0'),
        ('Supabase', 'Database + authentication', 'Free tier',    '$0'),
        ('Domain',   'Custom URL (not yet)',      '—',            '$0'),
        ('',         '',                          'TOTAL',        '$0'),
    ]
)

add_bullet(doc, "Current total cost: $0/month", bold_part="$0/month")
add_bullet(doc, "No credit card required at current usage level")

# ════════════════════════════════════════════════════════════
# 4. FREE TIER LIMITS
# ════════════════════════════════════════════════════════════
add_section(doc, "4. Free Tier Limits")

add_bullet(doc, "Supabase Free:", bold_part="Supabase Free:")
add_bullet(doc, "Database storage: 500MB", indent=1)
add_bullet(doc, "File storage: 1GB", indent=1)
add_bullet(doc, "Monthly active users: 50,000", indent=1)
add_bullet(doc, "Bandwidth: 5GB/month", indent=1)
add_bullet(doc, "Vercel Free:", bold_part="Vercel Free:")
add_bullet(doc, "Bandwidth: 100GB/month", indent=1)
add_bullet(doc, "Unlimited deployments", indent=1)
add_bullet(doc, "1 team member", indent=1)

# ════════════════════════════════════════════════════════════
# 5. STORAGE PROJECTIONS
# ════════════════════════════════════════════════════════════
add_section(doc, "5. Storage Projections")
add_note(doc, "Each patient record ≈ 5–10KB (116 fields + body map JSONB data)")

add_table(doc,
    headers=['Scenario', 'Patients/Month', 'Storage/Year', 'Free Tier Lasts'],
    rows=[
        ('1 hospital (pilot)',    '100–200',  '~12–24 MB',  'Indefinitely'),
        ('5 hospitals',           '~500',     '~60 MB',     '~8 years'),
        ('10 hospitals',          '~1,000',   '~120 MB',    '~4 years'),
        ('Large scale (20+ hosp)','~2,000+',  '~240 MB+',   '~2 years'),
    ]
)

# ════════════════════════════════════════════════════════════
# 6. COST PROJECTIONS
# ════════════════════════════════════════════════════════════
add_section(doc, "6. Cost Projections")

add_table(doc,
    headers=['Scenario', 'Vercel', 'Supabase', 'Total/Month'],
    rows=[
        ('1 hospital pilot today',              'Free',    'Free',     '$0'),
        ('1,000 patients/mo — years 1–4',       'Free',    'Free',     '$0'),
        ('1,000 patients/mo — after year 4',    'Free',    'Pro $25',  '$25'),
        ('5+ hospitals at scale',               'Pro $20', 'Pro $25',  '$45'),
    ]
)

add_bullet(doc, "Supabase Pro ($25/mo) adds: 8GB database, daily backups, no usage limits", bold_part="Supabase Pro")
add_bullet(doc, "Vercel Pro ($20/mo) adds: team collaboration, custom domains, advanced analytics", bold_part="Vercel Pro")

# ════════════════════════════════════════════════════════════
# 7. DOMAIN
# ════════════════════════════════════════════════════════════
add_section(doc, "7. Custom Domain (Optional)")

add_table(doc,
    headers=['Domain Option', 'Cost/Year', 'Recommended'],
    rows=[
        ('respondguatemala.org',  '~$15',    '✓ Best option'),
        ('traumaregistry.org',    '~$15',    ''),
        ('respond-gt.org',        '~$12',    ''),
        ('respondguatemala.com',  '~$12',    ''),
        ('.gt (Guatemala)',        '~$50–100','Not needed'),
    ]
)

add_bullet(doc, "Recommended: respondguatemala.org — clear, professional, appropriate for grants and publications", bold_part="respondguatemala.org")
add_bullet(doc, ".org extension is preferred for non-profit / research projects over .com")
add_bullet(doc, "Connecting domain to Vercel: free, takes ~5 minutes")
add_bullet(doc, "SSL certificate (https://): free, auto-generated by Vercel")
add_bullet(doc, "Total with domain: ~$15/year (~$1.25/month)", bold_part="~$15/year")

# ════════════════════════════════════════════════════════════
# 8. WHAT IS NEEDED TO LAUNCH AT 1 HOSPITAL
# ════════════════════════════════════════════════════════════
add_section(doc, "8. What Is Needed to Launch at 1 Hospital in Guatemala")

add_bullet(doc, "Technical setup (~1–2 hours):", bold_part="Technical setup (~1–2 hours):")
add_bullet(doc, "Add hospital name to the facility dropdown in the form", indent=1)
add_bullet(doc, "Create 1 admin account for the hospital", indent=1)
add_bullet(doc, "Create registrar accounts for data entry staff", indent=1)
add_bullet(doc, "Run 2–3 test records end-to-end before go-live", indent=1)

add_bullet(doc, "Operational:", bold_part="Operational:")
add_bullet(doc, "Staff training session: ~30–60 min (interface is straightforward)", indent=1)
add_bullet(doc, "Designate a record verifier (admin role)", indent=1)
add_bullet(doc, "Confirm devices: works on tablet, phone, or desktop", indent=1)

add_bullet(doc, "Regulatory:", bold_part="Regulatory:")
add_bullet(doc, "Research use → IRB / ethics committee approval required", indent=1)
add_bullet(doc, "Data privacy agreement with the hospital", indent=1)
add_bullet(doc, "Clinical/operational use only → likely no IRB needed", indent=1)

# ════════════════════════════════════════════════════════════
# 9. SCALABILITY
# ════════════════════════════════════════════════════════════
add_section(doc, "9. Scalability")
add_bullet(doc, "Architecture supports multiple hospitals simultaneously with no code changes")
add_bullet(doc, "Each hospital sees only its own data (row-level security in Supabase)")
add_bullet(doc, "Adding a new hospital = creating user accounts only (~5 min)")
add_bullet(doc, "Can scale to dozens of hospitals before hitting any meaningful cost threshold")
add_bullet(doc, "Offline-first design handles poor internet connectivity in low-resource settings")

# ════════════════════════════════════════════════════════════
# 10. BOTTOM LINE
# ════════════════════════════════════════════════════════════
add_section(doc, "10. Bottom Line")
add_bullet(doc, "The system is built, deployed, and working at $0/month", bold_part="$0/month")
add_bullet(doc, "A 1-hospital pilot can be launched in under 1 day of setup", bold_part="1 day of setup")
add_bullet(doc, "Cost only increases after year 4 at high volume, maxing out at $45/month at full scale", bold_part="$45/month")
add_bullet(doc, "With a custom domain: total annual cost is ~$15/year", bold_part="~$15/year")
add_bullet(doc, "The primary barrier to launch is not technology or cost — it is IRB approval if used for research", bold_part="IRB approval")

# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════
OUT = '/Users/dariorocha/Claude/trauma-registry/docs/RESPOND_Guatemala_Project_Report.docx'
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"✓ Saved: {OUT}")
